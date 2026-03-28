import pandas as pd
import os
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font, Alignment
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime

# 1. Configuración de archivos
excel_input = "calificaciones.xlsx"
word_folder = "cartas_padres"

# Crear carpeta para los documentos Word si no existe
if not os.path.exists(word_folder):
    os.makedirs(word_folder)

def obtener_fecha_espanol():
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    ahora = datetime.now()
    dia = ahora.day
    mes = meses[ahora.month]
    año = ahora.year
    return f"{dia} de {mes} de {año}"

def set_cell_shading(cell, fill_color):
    """
    Función para aplicar sombreado a una celda de una tabla en Word.
    fill_color debe ser un string hexadecimal (ej: 'C6EFCE')
    """
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def cargar_datos_excel():
    print(f"Cargando datos desde: {excel_input}...")
    if not os.path.exists(excel_input):
        print(f"Error: No se encontró el archivo {excel_input}")
        return None
    
    # Intentamos leerlo como CSV primero (por si acaso) y luego como Excel
    try:
        df = pd.read_csv(excel_input)
    except:
        df = pd.read_excel(excel_input)
    
    return df

def generar_cartas_word(df):
    print(f"Generando cartas en Word con tablas sombreadas en la carpeta: {word_folder}...")
    fecha_es = obtener_fecha_espanol()
    
    columnas = df.columns.tolist()
    materias = columnas[1:-1]
    
    for index, row in df.iterrows():
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Fecha en ESPAÑOL (Derecha)
        p_fecha = doc.add_paragraph()
        p_fecha.alignment = 2 
        p_fecha.add_run(f"Ciudad de México, a {fecha_es}")

        doc.add_paragraph("\nDIRIGIDA AL PADRE DE FAMILIA")
        doc.add_paragraph(f"Presente de: {row['Alumno']}\n")
        doc.add_paragraph("Estimado Padre de Familia:")

        cuerpo = (
            f"Por medio de la presente, hacemos de su conocimiento los resultados académicos "
            f"obtenidos por su hijo(a) {row['Alumno']} correspondientes al primer reporte del ciclo escolar. "
            f"A continuación se detallan las calificaciones por asignatura:"
        )
        doc.add_paragraph(cuerpo)

        # Tabla de calificaciones en Word
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Asignatura'
        hdr_cells[1].text = 'Calificación'
        
        # Color del encabezado de la tabla (Gris oscuro)
        set_cell_shading(hdr_cells[0], "333333")
        set_cell_shading(hdr_cells[1], "333333")
        # Texto blanco para el encabezado
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
        
        # Determinar color de sombreado según promedio
        try:
            promedio_val = float(row['Promedio'])
            # Verde claro: C6EFCE, Rojo claro: FFC7CE
            color_fondo = "C6EFCE" if promedio_val >= 6 else "FFC7CE"
        except:
            color_fondo = "FFFFFF" # Blanco por defecto

        for materia in materias:
            row_cells = table.add_row().cells
            row_cells[0].text = materia
            row_cells[1].text = str(row[materia])
            # Aplicar sombreado a las celdas de la fila
            set_cell_shading(row_cells[0], color_fondo)
            set_cell_shading(row_cells[1], color_fondo)
            
        # Promedio final (también sombreado en una fila extra opcional o párrafo)
        p_promedio = doc.add_paragraph()
        p_promedio.add_run(f"\nPROMEDIO GENERAL: {row['Promedio']}").bold = True

        if float(row['Promedio']) < 6:
            p_rec = doc.add_paragraph()
            p_rec.add_run("RECOMENDACIÓN: Se sugiere al alumno dedicar tiempo extra al estudio y solicitar asesorías en las materias con bajo desempeño para mejorar sus resultados en el próximo reporte.")

        doc.add_paragraph("\nSin más por el momento, agradecemos su atención y quedamos a su disposición para cualquier duda o aclaración.")
        p_firma = doc.add_paragraph()
        p_firma.alignment = 1 
        p_firma.add_run("\n\n__________________________\nAtentamente,\nLa Dirección Escolar")

        safe_name = "".join([c for c in str(row['Alumno']) if c.isalnum() or c==' ']).replace(' ', '_')
        file_name = f"Carta_{safe_name}.docx"
        doc.save(os.path.join(word_folder, file_name))

if __name__ == "__main__":
    df_datos = cargar_datos_excel()
    if df_datos is not None:
        generar_cartas_word(df_datos)
        print(f"\n¡Proceso completado! Tablas sombreadas en Word y fecha en español.")
